from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/keshavtyagi/Developer/WealthLens/output/cv/Keshav_Tyagi_CV.docx")

BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
LIGHT = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.7)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 15, ACCENT, 12, 5),
        ("Heading 2", 12.5, ACCENT, 8, 3),
        ("Heading 3", 11.5, BLUE, 5, 2),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Bullet 2"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.line_spacing = 1.10


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("Keshav Tyagi")
    set_run_font(run, size=24, bold=True, color=BLACK)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    c = contact.add_run(
        "(813) 327-9470 | Tampa, FL | keshav54@usf.edu | linkedin.com/in/keshav--tyagi"
    )
    set_run_font(c, size=10.5, color=GRAY)
    set_paragraph_border_bottom(contact)


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def add_role(doc, title, organization=None, location=None, dates=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    left = title
    if organization:
        left += f" | {organization}"
    if location:
        left += f" - {location}"
    r = p.add_run(left)
    set_run_font(r, size=11.2, bold=True, color=BLACK)
    if dates:
        d = p.add_run(f"    {dates}")
        set_run_font(d, size=10.5, italic=True, color=GRAY)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=BLACK)


def add_plain(doc, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=10.7, color=BLACK)
    return p


def add_skills_table(doc):
    rows = [
        ("Languages", "Python, TypeScript, JavaScript, C++, SQL"),
        ("Frontend", "Next.js, React, Tailwind CSS, ShadCN UI, Recharts, responsive UI/UX"),
        ("Backend & Data", "Node.js, Fastify, Express, REST APIs, PostgreSQL, Prisma ORM, Neon, Redis"),
        ("AI / ML", "FastAPI, Pandas, NumPy, scikit-learn, Gemini LLM, AI agents, Plaid API, forecasting, anomaly detection"),
        ("DevOps", "Git, GitHub Actions CI/CD, Vercel, Render, production debugging, testing"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    set_table_borders(table)
    widths = [Inches(1.35), Inches(5.35)]
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if j == 0:
                set_cell_shading(cell, LIGHT)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            set_run_font(run, size=10.2, bold=(j == 0), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build():
    doc = Document()
    style_doc(doc)
    add_header(doc)

    add_section_heading(doc, "Professional Profile")
    add_plain(
        doc,
        "Computer Science undergraduate at the University of South Florida with hands-on experience building full-stack, AI-enabled web applications. Project work spans financial technology, real-time public safety data, LLM-assisted user workflows, REST APIs, PostgreSQL data systems, and production deployment across modern cloud platforms.",
    )

    add_section_heading(doc, "Education")
    add_role(doc, "B.S. Computer Science", "University of South Florida", "Tampa, FL")

    add_section_heading(doc, "Technical Competencies")
    add_skills_table(doc)

    add_section_heading(doc, "Selected Technical Projects")
    add_role(doc, "WealthLens - AI Financial Copilot", "Full-Stack + AI Engineering", dates="2026")
    add_plain(
        doc,
        "Technologies: Next.js, React, TypeScript, Node.js/Fastify, PostgreSQL, Prisma, Python FastAPI, scikit-learn, Gemini LLM, Plaid API, Vercel",
        after=2,
    )
    add_bullets(
        doc,
        [
            "Built a production-deployed AI financial copilot with REST API endpoints, Clerk authentication, Plaid bank-account connectivity, real-time spending tracking, and personalized financial insights.",
            "Engineered a Python FastAPI analytics service using Pandas, NumPy, and scikit-learn to forecast cash flow, score financial health, and detect subscriptions and anomalies from transaction data.",
            "Integrated Gemini LLM with transaction-aware context for advisor chat, automated weekly reports, and goal planning; deployed across Vercel, Render, and Neon with GitHub Actions CI/CD.",
        ]
    )

    add_role(doc, "BayGuard Tampa - Disaster Intelligence Web App", "Hackathon", dates="2026")
    add_plain(
        doc,
        "Technologies: React 19, Vite, TypeScript, Node.js/Express, Google Maps API, Gemini AI, Redis, Vercel",
        after=2,
    )
    add_bullets(
        doc,
        [
            "Collaborated in a team of four to build a Tampa disaster intelligence app that pulls live data from NWS, NOAA, NHC, FL-511, and Tampa Electric into a real-time neighborhood risk map.",
            "Architected a multi-agent AI layer using Google Gemini, with specialized weather, flood, storm, and report-verification agents feeding a final synthesis layer for readable evacuation guidance and situational summaries.",
            "Built resident report submission with AI-assisted verification, evidence cross-referencing, SMS alerting, subscriber management, cooldown controls, and simulation modes for flood and hurricane drills.",
        ]
    )

    doc.add_page_break()
    add_section_heading(doc, "Professional Experience")
    add_role(doc, "Student Assistant", "USF Office of the Provost (Academic Affairs)", "Tampa, FL", "Jun 2026 - Present")
    add_bullets(
        doc,
        [
            "Manage and triage JIRA tickets, coordinate task workflows, and collaborate with office staff to resolve administrative and technical requests efficiently.",
            "Use Office 365 tools including Teams, Word, and Excel to support communication, documentation, and daily office operations.",
        ]
    )

    add_role(doc, "Import & Export Assistant", "Family Business", "Riverview, FL", "2022 - 2024")
    add_bullets(
        doc,
        [
            "Managed vendor pricing research, bulk inventory procurement, shipment coordination, and online inventory tracking.",
        ]
    )

    add_role(doc, "Cashier & Stock Associate", "Local Grocery Store", "Gibsonton, FL", "2023 - 2024")
    add_bullets(
        doc,
        [
            "Processed high-volume transactions, restocked inventory, and maintained customer-facing areas in a fast-paced retail environment.",
        ]
    )

    add_section_heading(doc, "Community Service & Activities")
    add_role(doc, "Volunteer", "Feeding Tampa Bay", "Tampa, FL", "2022 - 2023")
    add_bullets(
        doc,
        [
            "Supported high-volume community food distribution through meal preparation, pantry operations, and service logistics.",
        ]
    )
    add_role(doc, "JV Wrestling", dates="Student Activity")
    add_bullets(
        doc,
        [
            "Competed in tournaments and developed discipline, teamwork, and resilience through consistent training.",
        ]
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(4)
        run = footer.add_run("Keshav Tyagi - Curriculum Vitae")
        set_run_font(run, size=8.5, color=GRAY)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
